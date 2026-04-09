"""
Generate Protocol 6 Results Paper as .docx
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
section = doc.sections[0]
section.top_margin    = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin   = Inches(1.25)
section.right_margin  = Inches(1.25)

# ── Style helpers ─────────────────────────────────────────────────────────────
normal = doc.styles['Normal']
normal.font.name = 'Times New Roman'
normal.font.size = Pt(12)

def set_font(run, bold=False, italic=False, size=12, name='Times New Roman'):
    run.bold   = bold
    run.italic = italic
    run.font.name = name
    run.font.size = Pt(size)

def heading(text, level=1):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13 if level == 1 else 12)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, indent=False, space_after=6):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if indent:
        p.paragraph_format.first_line_indent = Inches(0.3)
    p.paragraph_format.space_after = Pt(space_after)
    return p

def italic_body(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def mixed_body(*parts):
    """parts = list of (text, bold, italic) tuples"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    for text, bold, italic in parts:
        run = p.add_run(text)
        run.bold   = bold
        run.italic = italic
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_table_row(table, cells, bold_first=False, shaded=False):
    row = table.add_row()
    for i, (cell, text) in enumerate(zip(row.cells, cells)):
        cell.text = ''
        run = cell.paragraphs[0].add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        if bold_first and i == 0:
            run.bold = True
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

# ── TITLE BLOCK ───────────────────────────────────────────────────────────────
title_p = doc.add_paragraph()
title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = title_p.add_run(
    "Emergent Constraint Fields Are Causally Active But Do Not Outperform "
    "Fixed External Rules: A Preregistered Null on Passive Emergence as a "
    "Governance Strategy"
)
title_run.bold = True
title_run.font.name = 'Times New Roman'
title_run.font.size = Pt(14)
title_p.paragraph_format.space_after = Pt(12)

# Author / affiliation block
for line in [
    "Bruce Tisler",
    "Quantum Inquiry (quantuminquiry.org)",
    "ORCID: 0009-0009-6344-5334",
    "Protocol 6 of the Ethics as Emergent Constraint Response Series",
    "Preregistration DOI: 10.5281/zenodo.19297509",
    "Repository: btisler-DS/emergent-constraint-marl",
    "Date: 2026-04-09",
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(line)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    p.paragraph_format.space_after = Pt(2)

doc.add_paragraph()

# ── ABSTRACT ──────────────────────────────────────────────────────────────────
heading("Abstract")
body(
    "Does the origin of a constraint — emergent from agent activity versus imposed "
    "from outside — determine whether that constraint produces alignment-relevant "
    "behavioral outcomes in multi-agent systems? Protocol 6 of the Ethics as "
    "Emergent Constraint Response series tested this question using a preregistered "
    "four-condition design. Three-agent reinforcement learning systems were trained "
    "under: Condition A (emergent constraint field, local perception), Condition B "
    "(emergent constraint field, global perception), Condition C (fixed external "
    "constraint, matched cost), and Condition D (no constraint). Two hundred "
    "confirmatory runs (50 seeds per condition, 500 epochs each) were executed "
    "against a Zenodo-preregistered protocol (SHA-256: "
    "c286a89037966e56630f5d3ce4cdb4a1621ce7ce9a0f7e7d477a06b8495dffc3)."
)
body(
    "Results were mixed. The mechanistic prediction was robustly confirmed: within-run "
    "correlation between field entropy and behavioral structure (sustained_structure_score "
    "proxy) yielded a median r = \u22120.680 in Condition A (Wilcoxon W = 3.0, p < 0.001), "
    "closely replicating the pilot finding. The emergent constraint field is causally "
    "active and creates a temporal coupling signature entirely absent in the fixed-"
    "constraint and unconstrained conditions. However, the primary behavioral claim was "
    "not confirmed: Condition A did not produce significantly higher reward than Condition "
    "C (p = 0.069), and the behavioral homogenization attractor documented across "
    "Protocols 2\u20135 persisted in all constrained conditions (mean query_rate \u2248 0.78 "
    "in A, B, and C). The H3 reversal \u2014 global field perception produced more "
    "behavioral variance than local, contrary to all committee predictions \u2014 was "
    "the largest empirical surprise. The governance implication is direct: an emergent "
    "constraint field that co-constitutes itself through agent signaling behaves "
    "mechanistically differently from an imposed rule, but this mechanistic difference "
    "does not translate to better alignment outcomes in this architecture. Passive "
    "emergence is insufficient as a governance strategy."
)

# ── 1. INTRODUCTION ───────────────────────────────────────────────────────────
heading("1. Introduction")
body(
    "The Ethics as Emergent Constraint Response series began with a deceptively simple "
    "observation: regulatory constraints imposed on artificial agents do not produce the "
    "behavioral outcomes they nominally target. Protocol 2 (Tisler, 2026a) demonstrated "
    "that agents subject to a fixed ethical tax learn to satisfy the tax through query "
    "flooding — a compliance-satisfying but alignment-defeating attractor with effect "
    "size d = \u22122.18. Protocol 3 (Tisler, 2026b) showed that epistemic opacity "
    "amplifies this failure: when agents cannot observe the enforcement mechanism, "
    "gaming increases (H1 inverted, d = +2.22). Protocol 4 (Tisler, 2026c) added "
    "recursive self-modeling depth and found that self-inclusion and ethical output "
    "are behaviorally dissociated \u2014 deeper self-models increase sacrifice capacity "
    "without increasing alignment. Protocol 5 (Tisler, 2026d) produced a complete "
    "null: temporal integration span and prosocial constraint architecture are jointly "
    "insufficient to rescue alignment. All five primary hypotheses failed."
)
body(
    "These results converge on a pattern: constraint architectures that are imposed "
    "from outside the system do not produce the structural conditions for alignment in "
    "this three-agent configuration, regardless of their form, opacity, or the "
    "cognitive depth of the agents subject to them. Protocol 6 asks a question that "
    "the prior protocols could not address: what if the constraint is not imposed but "
    "emerges? If agents co-constitute the constraint field through their own signal "
    "production, does that change the outcome?"
)
body(
    "The motivation is both theoretical and practical. In governance literature, "
    "there is a recurring argument that norms which emerge from within a community "
    "are more stable and more effective than rules enforced from outside it "
    "(Ostrom, 1990; Axelrod, 1986). The slime mold Dictyostelium discoideum provides "
    "a biological analogue: individual cells emit cyclic AMP signals that propagate "
    "as a field, and this self-assembled field coordinates collective behavior in ways "
    "that no single cell plans. Protocol 6 tests whether an analogous mechanism "
    "\u2014 agents emitting signals that diffuse into a shared constraint landscape "
    "they then navigate \u2014 produces qualitatively different outcomes than an "
    "externally imposed ethical tax of equivalent cost."
)
body(
    "The answer is mechanistically yes and behaviorally no. This paper reports the "
    "full confirmatory results, the statistical evidence for both findings, the "
    "performance of a five-model AI committee that predicted outcomes in advance, "
    "and the implications for agent governance system design."
)

# ── 2. THEORETICAL BACKGROUND ─────────────────────────────────────────────────
heading("2. Theoretical Background")
body(
    "The constraint field mechanism introduced in Protocol 6 is inspired by the "
    "Dictyostelium aggregation model of emergent coordination (Goldbeter, 1996; "
    "Kessin, 2001). In Dictyostelium, individual amoebae under starvation stress "
    "emit cyclic AMP pulses. These pulses diffuse, decay, and are relayed by "
    "neighboring cells, producing a propagating wave that serves as a collective "
    "coordination signal. No cell plans the wave. The wave is the aggregate "
    "consequence of individual cell behavior responding to local signal density. "
    "The resulting structure \u2014 a fruiting body with differentiated stalk and "
    "spore cells \u2014 is a behavioral outcome that neither individual cells nor "
    "the local signal field specifies."
)
body(
    "The Protocol 6 constraint field F (length 3, one scalar per agent position) "
    "operates on the same structural logic. Agents emit communication signals "
    "(DECLARE, QUERY, RESPOND) whose weighted sum increments each agent's field "
    "position. The field then diffuses across positions (diffusion coefficient "
    "dc = 0.1), decays each epoch (decay rate dr = 0.05), and returns to each "
    "agent as an additional observation dimension. Each agent perceives only its "
    "own field value (Condition A: local perception) or the full field vector "
    "(Condition B: global perception). The effective cost of communication is "
    "modulated by the local field value: agents in high-field positions pay higher "
    "costs for additional signals."
)
body(
    "Two distinct claims are tested. The mechanistic claim holds that the field "
    "creates a temporal coupling between field structure and behavioral outcomes: "
    "when field entropy is high (field is uniformly distributed), behavioral "
    "structure (sustained coordination) should be lower, because no agent faces "
    "elevated cost pressure that would break the query-flooding equilibrium. When "
    "field entropy is low (field is heterogeneous), some agents are under elevated "
    "cost pressure, potentially disrupting the homogenization attractor. This "
    "prediction is operationalized as a negative within-run Spearman correlation "
    "between field entropy and sustained_structure_score."
)
body(
    "The behavioral claim is stronger: the emergent constraint should not only "
    "create a mechanistic signature but should produce better behavioral outcomes "
    "than a well-calibrated fixed external tax of equivalent mean cost. This is "
    "the governance-relevant claim. If true, it suggests that the origin of a "
    "constraint \u2014 whether it grows from within the agent community or is "
    "imposed from outside \u2014 determines alignment outcomes. If false, the "
    "mechanism is real but inconsequential for the properties that matter in "
    "governance contexts."
)
body(
    "The behavioral homogenization attractor documented across Protocols 2\u20135 "
    "provides the null baseline. In this architecture, agents consistently converge "
    "to high, uniform query rates regardless of constraint design. The question "
    "Protocol 6 poses is whether an internally co-constituted constraint can break "
    "this attractor in ways that externally imposed constraints cannot."
)

# ── 3. METHODS ────────────────────────────────────────────────────────────────
heading("3. Methods")

heading("3.1 Experimental Design", level=2)
body(
    "Protocol 6 employed a four-condition between-seeds design. All conditions used "
    "the same three-agent MARL harness (AgentA: RNN, AgentB: CNN, AgentC: GNN) "
    "ported from the Protocol 5 codebase with modifications documented below. "
    "Each condition ran 50 independent seeds (total: 200 runs), 500 epochs each."
)

# Conditions table
doc.add_paragraph()
tbl = doc.add_table(rows=1, cols=3)
tbl.style = 'Table Grid'
hdr = tbl.rows[0].cells
for cell, text in zip(hdr, ["Condition", "Description", "Field"]):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)

for row_data in [
    ("A — Emergent Local",   "Field active; AgentA observes own field value only (obs_dim=13)",  "Active (dc=0.1, dr=0.05)"),
    ("B — Emergent Global",  "Field active; AgentA observes full field vector (obs_dim=15)",       "Active (dc=0.1, dr=0.05)"),
    ("C — Fixed External",   "No field; fixed communication tax = 1.2287 per signal (matched cost)", "Inactive"),
    ("D — No Constraint",    "No field; no communication cost; unconstrained baseline",              "Inactive"),
]:
    add_table_row(tbl, row_data, bold_first=True)

doc.add_paragraph()

heading("3.2 Parameters", level=2)
body(
    "All parameters were preregistered prior to confirmatory execution. The grid "
    "environment: grid_size=20, z_layers=8, energy_budget=100, max_steps=64. "
    "Field parameters (Conditions A and B): diffusion_coefficient=0.1, "
    "decay_rate=0.05, selected from the 60-run pilot parameter sweep "
    "(6 combinations \u00d7 10 seeds). Signal weights (Deviation 1, see below): "
    "DECLARE=0.03, QUERY=0.01, RESPOND=0.02. Fixed cost multiplier (Condition C): "
    "1.2287, derived from the pilot mean field cost to match expected constraint "
    "overhead across conditions. Infrastructure: SimplePod 6\u00d7RTX 3060, "
    "breathtaking-porcupine, EU-PL-01."
)

heading("3.3 Logged Deviations", level=2)
body(
    "Two deviations from the initial design are on record, both preregistered "
    "before confirmatory execution."
)
mixed_body(
    ("Deviation 1 (commit adca6fc, pre-pilot): ", True, False),
    ("Signal weights were scaled 0.1\u00d7 from originally proposed values "
     "(DECLARE 0.3\u21920.03, QUERY 0.1\u21920.01, RESPOND 0.2\u21920.02) to prevent "
     "field saturation identified in preliminary analysis. The preregistration "
     "reflects the scaled values.", False, False)
)
mixed_body(
    ("Deviation 2 (post-pilot, 2026-03-28): ", True, False),
    ("The field formation criterion was revised from field_std > 0.05 sustained "
     "for more than 50 consecutive epochs to an activity-based criterion "
     "(field_mean > 0.01 after epoch 50). Concurrently, the primary research "
     "question was revised from spatial differentiation of the field to temporal "
     "coupling between field structure and behavioral outcomes. This revision was "
     "motivated by the pilot finding that behavioral homogenization collapses "
     "spatial field variance in all seeds, making the spatial framing of H1 and "
     "H3 empirically unresolvable. The SHA-256 hash of the final preregistration "
     "PDF (v3) was locked before confirmatory runs began.", False, False)
)

heading("3.4 Metric Availability Limitation", level=2)
body(
    "The preregistered primary metrics \u2014 sustained_structure_score (SSS) per seed "
    "and exploitation_loop_rate (ELR) per seed as scalar summaries \u2014 were not "
    "persisted as per-seed outputs in the initial confirmatory run. The epoch-level "
    "time series were computed during runs but not written to disk. A post-run fix "
    "(commit b1c1bb1) corrected persistence for future runs, but the 200 confirmatory "
    "seeds were already complete. All hypothesis tests reported below use the closest "
    "available proxies: entropy_sss_correlation (within-run Pearson correlation "
    "between field_entropy and SSS time series) for H1 and the mechanistic prediction, "
    "and final_avg_reward_A for H2. This limitation is acknowledged throughout "
    "the Results section. No hypothesis test is presented as confirmatory of the "
    "preregistered specification where the proxy departs from it."
)

# ── 4. RESULTS ────────────────────────────────────────────────────────────────
heading("4. Results")

heading("4.1 Descriptive Statistics", level=2)
body("Table 1 reports summary statistics across all 50 seeds per condition.")

doc.add_paragraph()
tbl2 = doc.add_table(rows=1, cols=6)
tbl2.style = 'Table Grid'
headers = ["Condition", "Mean Reward\n(\u00b1SD)", "Mean Query\nRate (\u00b1SD)",
           "Entropy-SSS\nCorr (\u00b1SD)", "Field\nMean", "Field\nCollapsed"]
for cell, text in zip(tbl2.rows[0].cells, headers):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)

rows_data = [
    ("A (emerg. local)",  "30.894 (\u00b15.134)", "0.782 (\u00b10.167)", "\u22120.639 (\u00b10.233)", "0.238", "0/50"),
    ("B (emerg. global)", "26.784 (\u00b18.406)", "0.770 (\u00b10.205)", "\u22120.582 (\u00b10.283)", "0.239", "0/50"),
    ("C (fixed ext.)",    "29.540 (\u00b15.208)", "0.784 (\u00b10.188)", "0.000 (\u00b10.000)",        "0.000", "50/50"),
    ("D (no constr.)",    "38.929 (\u00b14.065)", "0.305 (\u00b10.237)", "0.000 (\u00b10.000)",        "0.000", "50/50"),
]
for rd in rows_data:
    add_table_row(tbl2, rd, bold_first=True)

p_caption = doc.add_paragraph()
r = p_caption.add_run("Table 1. Descriptive statistics (n = 50 per condition, 500 epochs per run).")
r.italic = True
r.font.name = 'Times New Roman'
r.font.size = Pt(10)
p_caption.paragraph_format.space_after = Pt(8)

body(
    "Several patterns are notable before hypothesis testing. First, the field is "
    "functioning in A and B: both show non-zero mean field values (~0.24) and "
    "zero collapses across all 50 seeds. Second, behavioral homogenization "
    "dominates A, B, and C: query_rate converges to approximately 0.78 in all "
    "three constrained conditions. Condition D is the exception at query_rate = "
    "0.305, consistent with unconstrained agents learning that silence is optimal "
    "when signaling is costly but no offsetting reward gradient exists. Third, "
    "entropy-SSS correlation is substantially negative in A (\u22120.639) and B "
    "(\u22120.582) and exactly zero in C and D (no field, no coupling). The "
    "qualitative difference between emergent and non-emergent conditions is "
    "apparent in this metric. Fourth, reward ordering is D > A > C > B: "
    "unconstrained agents achieve highest reward; global field perception (B) "
    "produces lowest reward with highest inter-seed variance (SD = 8.406), "
    "suggesting less stable behavioral attractors under global field information."
)

heading("4.2 H1 \u2014 Interrogative Diversity (Proxy Result)", level=2)
body(
    "Preregistered: Condition A shows higher sustained_structure_score than "
    "Condition C (Mann-Whitney U, one-tailed, p < 0.05, Cohen\u2019s d > 0.5)."
)
body(
    "Proxy used: entropy_sss_correlation (within-run Pearson correlation between "
    "field_entropy and SSS time series). Direction tested: A more negative than C "
    "(stronger field-behavior coupling in the emergent condition)."
)
mixed_body(
    ("Result: ", True, False),
    ("Mann-Whitney U = 50.0, p < 0.001 (one-tailed), rank-biserial r = 0.96. "
     "Mean A = \u22120.639, Mean C = 0.000. ", False, False),
    ("PROXY SUPPORTED.", True, False)
)
body(
    "The effect size (r = 0.96) is near-ceiling, reflecting near-complete "
    "separation: 49 of 50 Condition A seeds show negative entropy_sss_correlation, "
    "while all 50 Condition C seeds show exactly 0.0 by construction (no field "
    "mechanism, no coupling). This confirms that the emergent field creates "
    "qualitatively distinct within-run behavioral dynamics. However, this is not "
    "a direct test of mean SSS level as preregistered, and this distinction must "
    "be acknowledged. The proxy establishes that the field mechanism is causally "
    "active; it does not establish that the resulting SSS is higher."
)

heading("4.3 H2 \u2014 Constraint Gaming Suppression (Proxy Result)", level=2)
body(
    "Preregistered: Condition A shows lower exploitation_loop_rate than both "
    "Condition C and Condition D (two Mann-Whitney U tests, p < 0.025 each, "
    "Bonferroni-corrected). Both comparisons must succeed."
)
body(
    "Proxy used: final_avg_reward_A (inverted proxy: lower reward proxies higher "
    "gaming cost if gaming = wasteful communication overhead)."
)
mixed_body(
    ("H2a (A vs. D): ", True, False),
    ("Mann-Whitney U = 245.0, p < 0.001, r = 0.804. Mean A = 30.894, "
     "Mean D = 38.929. ", False, False),
    ("SUPPORTED at Bonferroni threshold. ", True, False),
    ("The unconstrained condition achieves substantially higher reward, "
     "confirming that the emergent constraint field imposes coordination "
     "overhead not offset by reward gains.", False, False)
)
mixed_body(
    ("H2b (A vs. C): ", True, False),
    ("Mann-Whitney U = 1,466.0, p = 0.069, r = \u22120.173. Mean A = 30.894, "
     "Mean C = 29.540. ", False, False),
    ("NOT SUPPORTED (p > 0.025 Bonferroni threshold). ", True, False),
    ("The 1.35-point reward difference is in the predicted direction "
     "(emergent marginally more efficient) but does not reach significance. "
     "H2 as a whole fails: both comparisons are required, and H2b does not "
     "clear the corrected threshold.", False, False)
)
body(
    "Under the preregistered falsification criteria, H2\u2019s failure is the "
    "primary null finding of Protocol 6. The emergent constraint does not "
    "produce significantly better reward outcomes than a well-calibrated fixed "
    "external tax."
)

heading("4.4 H3 \u2014 Locality as Necessary Condition (Proxy Result)", level=2)
body(
    "Preregistered: Variance of within-run per-agent query_rate SD is greater "
    "in Condition A than Condition B (Levene\u2019s test, p < 0.05). Local field "
    "perception predicted to produce more behavioral heterogeneity than global."
)
body(
    "Proxy used: cross-seed variance of final_query_rate (captures between-seed "
    "training variability, not within-run agent-level divergence as preregistered)."
)
mixed_body(
    ("Result: ", True, False),
    ("Levene W = 1.44, p = 0.232. Var(A) = 0.0278, Var(B) = 0.0422 "
     "(ratio = 0.659). Direction reversed: Var(B) > Var(A). ", False, False),
    ("NOT SUPPORTED; direction inverted.", True, False)
)
body(
    "Global field perception (B) produces more between-seed behavioral variability "
    "than local field perception (A). This is opposite to the preregistered "
    "prediction and was not anticipated by any member of the AI committee. The "
    "reversal suggests that global field information destabilizes behavioral "
    "convergence across seeds \u2014 agents with access to the full field vector "
    "find a wider range of behavioral equilibria across training runs \u2014 while "
    "local perception constrains agents to a narrower, more consistent attractor "
    "basin. The within-run heterogeneity that H3 targets may still exist; the "
    "proxy metric cannot resolve this."
)

heading("4.5 Mechanistic Prediction \u2014 Temporal Coupling", level=2)
body(
    "Preregistered: Pooled Spearman correlation between per-epoch field_entropy "
    "and exploitation_loop_rate across Condition A seeds is negative (r < 0, "
    "p < 0.05). Secondary: per-seed Spearman correlation sign test, median r < 0."
)
body(
    "Proxy used: per-seed within-run Pearson correlation between field_entropy "
    "and SSS time series (entropy_sss_correlation). SSS and ELR are inversely "
    "related conceptually, so the predicted sign is preserved. Pearson rather "
    "than Spearman; sign direction intact."
)
mixed_body(
    ("Result: ", True, False),
    ("One-sample Wilcoxon signed-rank vs. 0, one-tailed: W = 3.0, p < 0.001. "
     "Median = \u22120.680, Mean = \u22120.639 (n = 50). ", False, False),
    ("STRONGLY SUPPORTED.", True, False)
)
body(
    "49 of 50 Condition A seeds show negative entropy_sss_correlation. The "
    "median (\u22120.680) closely replicates the pilot finding (r = \u22120.650) "
    "at the selected parameter combination (dc = 0.1, dr = 0.05), despite "
    "2.5\u00d7 longer runs (500 vs 200 epochs). The signal did not attenuate "
    "with extended training. When the constraint field is more uniformly "
    "distributed (high entropy, every agent under similar cost pressure), "
    "behavioral structure is lower. When the field is structured (some agents "
    "under elevated cost pressure), behavioral structure is higher. The field "
    "is doing causal work."
)

heading("4.6 H4 \u2014 Behavioral Differentiation (Exploratory)", level=2)
body(
    "Preregistered: behavioral_differentiation_index exceeds low-pressure baseline "
    "by 2 SD in \u226560% of Condition A seeds. Proxy: cross-seed std of final "
    "query_rate as differentiation index."
)
mixed_body(
    ("Result: ", True, False),
    ("std ordering: D = 0.237 > B = 0.205 > C = 0.188 > A = 0.167. "
     "Predicted ordering: A > B > C \u2248 D. Direction fully reversed for A. ", False, False),
    ("NOT SUPPORTED.", True, False)
)
body(
    "Condition A shows the lowest behavioral differentiation of all four conditions "
    "on this proxy. The emergent local constraint produces the most homogeneous "
    "training outcomes, not the most differentiated. Combined with H3\u2019s "
    "reversal, the picture is consistent: local field perception acts as a "
    "behavioral stabilizer, converging agents toward a common attractor more "
    "reliably than other conditions."
)

heading("4.7 Summary of Hypothesis Tests", level=2)

doc.add_paragraph()
tbl3 = doc.add_table(rows=1, cols=4)
tbl3.style = 'Table Grid'
for cell, text in zip(tbl3.rows[0].cells, ["Hypothesis", "Test / Statistic", "Result", "Verdict"]):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)

hyp_rows = [
    ("H1 (proxy): A more negative entropy-SSS corr than C",
     "Mann-Whitney U = 50.0\np < 0.001, r = 0.96",
     "Mean A = \u22120.639, Mean C = 0.000",
     "SUPPORTED (proxy)"),
    ("H2a: A < D on reward",
     "Mann-Whitney U = 245.0\np < 0.001, r = 0.804",
     "Mean A = 30.89, Mean D = 38.93",
     "SUPPORTED"),
    ("H2b: A > C on reward",
     "Mann-Whitney U = 1,466\np = 0.069, r = \u22120.173",
     "Mean A = 30.89, Mean C = 29.54",
     "NOT SUPPORTED"),
    ("H3 (proxy): Var(A) > Var(B) on query_rate",
     "Levene W = 1.44\np = 0.232",
     "Var(A) = 0.028, Var(B) = 0.042; reversed",
     "NOT SUPPORTED"),
    ("Mechanistic: entropy-SSS corr < 0 in Cond. A",
     "Wilcoxon W = 3.0\np < 0.001",
     "Median = \u22120.680, Mean = \u22120.639",
     "STRONGLY SUPPORTED"),
    ("H4: BDI ordering A > B > C \u2248 D",
     "Descriptive\n(std ordering)",
     "D > B > C > A; reversed",
     "NOT SUPPORTED"),
]
for rd in hyp_rows:
    add_table_row(tbl3, rd, bold_first=False)

p_cap2 = doc.add_paragraph()
rc = p_cap2.add_run("Table 2. Hypothesis test summary. 3/6 tests supported at \u03b1 = 0.05.")
rc.italic = True
rc.font.name = 'Times New Roman'
rc.font.size = Pt(10)
p_cap2.paragraph_format.space_after = Pt(8)

# ── 5. AI COMMITTEE PREDICTIONS ───────────────────────────────────────────────
heading("5. AI Committee Predictions")
body(
    "Prior to confirmatory execution, five AI models \u2014 ChatGPT, Gemini, "
    "Grok, DeepSeek, and Claude (experiment manager) \u2014 committed predictions "
    "on all hypotheses. Predictions were logged in commit 7625fe1, before any "
    "confirmatory data existed. Scoring used the following key: CORRECT (direction "
    "and threshold match), PARTIAL (direction correct, magnitude or threshold "
    "missed), INCORRECT (does not match), N/A (cannot evaluate with available data)."
)

# Committee scorecard table
doc.add_paragraph()
tbl4 = doc.add_table(rows=1, cols=7)
tbl4.style = 'Table Grid'
for cell, text in zip(tbl4.rows[0].cells,
                      ["", "H1", "H2\nOutcome", "H3", "Mechanistic", "H4", "Score"]):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)

committee_rows = [
    ("ChatGPT",     "INCORRECT", "CORRECT",   "CORRECT",   "PARTIAL",  "CORRECT", "3.5/5"),
    ("Gemini",      "CORRECT",   "CORRECT*",  "INCORRECT", "CORRECT",  "N/A",     "3.5/4"),
    ("Grok",        "INCORRECT", "CORRECT",   "CORRECT",   "INCORRECT","CORRECT", "3.0/5"),
    ("DeepSeek",    "CORRECT",   "INCORRECT", "INCORRECT", "CORRECT",  "CORRECT", "3.0/5"),
    ("Claude (Mgr)","CORRECT",   "PARTIAL",   "INCORRECT", "CORRECT",  "CORRECT", "3.5/5"),
]
for rd in committee_rows:
    add_table_row(tbl4, rd, bold_first=True)

p_cap3 = doc.add_paragraph()
rc3 = p_cap3.add_run(
    "Table 3. AI committee prediction scorecard. *Gemini scored on A < D component; full H2 not explicit."
)
rc3.italic = True
rc3.font.name = 'Times New Roman'
rc3.font.size = Pt(10)
p_cap3.paragraph_format.space_after = Pt(6)

body(
    "Three findings stand out. First, Gemini produced the most precise quantitative "
    "prediction in the committee: the mechanistic effect would fall in the range "
    "r = \u22120.45 to \u22120.80. The actual median (\u22120.680) and mean (\u22120.639) "
    "both fall squarely within this range. This required correctly predicting that "
    "the pilot signal (r = \u22120.650 at 200 epochs) would hold at full strength "
    "through 500-epoch confirmatory runs, against Grok\u2019s theoretically "
    "sophisticated attenuation argument."
)
body(
    "Second, Grok produced the most significant miss in the committee: the "
    "mechanistic prediction of r \u2265 \u22120.15 (effectively null). Grok argued "
    "that agents would learn to optimize around the field in extended runs, making "
    "low-entropy field states reflect stable low-cost equilibria that enable "
    "rather than suppress exploitation. This mechanism did not materialize. The "
    "signal strengthened, not attenuated."
)
body(
    "Third, the H3 reversal \u2014 global field perception producing more behavioral "
    "variance than local \u2014 was predicted by no committee member. Every member "
    "who addressed the direction of H3 predicted Var(A) \u2265 Var(B). The actual "
    "result (Var(B)/Var(A) ratio = 1.52) was outside all five members\u2019 models "
    "of the system. This is the genuine empirical surprise of Protocol 6 and the "
    "finding that most directly challenges existing theoretical framing."
)

# ── 6. DISCUSSION ─────────────────────────────────────────────────────────────
heading("6. Discussion")

heading("6.1 What the Mechanistic Confirmation Means", level=2)
body(
    "The robust confirmation of the mechanistic prediction (median r = \u22120.680, "
    "W = 3.0, p < 0.001, 49/50 seeds negative) establishes that the emergent "
    "constraint field is causally active. This is not a null result dressed up "
    "as a finding. The within-run temporal coupling between field entropy and "
    "behavioral structure is a genuine phenomenon: when the constraint field is "
    "structured (low entropy, heterogeneous cost distribution), behavioral "
    "coordination is higher. When the field is uniform (high entropy, homogeneous "
    "costs), coordination is lower. The field is doing work that a fixed "
    "external tax cannot do, because a fixed tax has no entropy."
)
body(
    "This is also the first finding in the P2\u2013P6 series where the constraint "
    "mechanism produces a distinct, positive behavioral signature. Protocols 2\u20135 "
    "produced gaming, amplification, dissociation, and null results respectively. "
    "Protocol 6 produces a causally active mechanism. The Dictyostelium analogy "
    "holds at the structural level: agents co-constitute a field through their "
    "signals, and that field modulates their subsequent behavior. This is "
    "mechanistically different from externally imposed constraint."
)

heading("6.2 Why the Behavioral Null Is the Headline for Governance", level=2)
body(
    "The H2b failure (A vs. C, p = 0.069) is the governance-relevant finding. "
    "An emergent constraint field that demonstrably does causal work does not "
    "produce meaningfully better behavioral outcomes than a fixed external rule "
    "of equivalent average cost. The 1.35-point reward advantage of A over C "
    "is in the predicted direction but falls short of significance. For governance "
    "purposes, this means that the additional complexity of allowing constraints "
    "to emerge from within an agent community \u2014 the co-constitution of the "
    "constraint landscape through agent activity \u2014 does not pay off in the "
    "behavioral outcomes that matter."
)
body(
    "The deeper problem is H2a (A vs. D, r = 0.804): the unconstrained condition "
    "achieves substantially higher reward (38.929 vs. 30.894). Condition D agents "
    "discover that silence is optimal when signaling is costly but coordination "
    "gains do not offset those costs (query_rate = 0.305). The emergent constraint "
    "field prevents this silence strategy \u2014 agents in A must continue paying "
    "communication costs to maintain field structure \u2014 but the maintained "
    "field does not produce commensurate reward returns. Constraint creates cost "
    "without proportional benefit. This pattern is consistent across P2\u2013P6 "
    "regardless of constraint architecture."
)

heading("6.3 The H3 Reversal and Sandbox Design Implications", level=2)
body(
    "The H3 reversal \u2014 global field perception producing more behavioral "
    "variability than local \u2014 has direct implications for surveillance-based "
    "sandbox design. The intuition behind H3 was that agents with local information "
    "would develop more heterogeneous behavioral trajectories because each agent "
    "faces a different effective cost landscape. Full field visibility was expected "
    "to homogenize behavior by giving all agents identical information."
)
body(
    "The data contradict this. Global field access (Condition B) produces higher "
    "cross-seed variance in terminal query_rate (Var = 0.042) and higher reward "
    "variance (SD = 8.406 vs. 5.134 in A). Agents with full field information find "
    "more diverse behavioral equilibria across training seeds. The plausible "
    "mechanism: global field visibility introduces an additional information "
    "channel that agents can use in multiple ways, expanding the effective strategy "
    "space. Local perception forces a more constrained agent-field interaction that "
    "funnels agents toward a common attractor."
)
body(
    "For surveillance sandbox design, this implies that providing agents with more "
    "complete visibility into a shared constraint field does not guarantee behavioral "
    "convergence. Full transparency about the constraint landscape may actually "
    "increase behavioral diversity across training runs, complicating evaluation "
    "and certification. This is a non-obvious finding that challenges common "
    "assumptions in AI safety sandbox design."
)

heading("6.4 The SSS/ELR Metric Gap", level=2)
body(
    "The failure to persist sustained_structure_score and exploitation_loop_rate "
    "as per-seed scalar summaries is a genuine limitation. All confirmatory "
    "hypothesis tests use proxy measures that depart from the preregistered "
    "specifications in important ways. Commit b1c1bb1 corrects this for future "
    "runs, but cannot retroactively recover the primary metrics for the 200 "
    "completed seeds."
)
body(
    "The effect on conclusions is bounded but real. H1\u2019s proxy result "
    "(r = 0.96 on entropy-SSS correlation) is very likely conservative: the "
    "coupling coefficient captures whether the field creates temporal dynamics "
    "in SSS, not whether mean SSS is elevated. The primary H1 test (mean SSS "
    "higher in A than C) could in principle show a different pattern from the "
    "coupling coefficient, though this is unlikely given the mechanistic "
    "confirmation\u2019s strength. H2\u2019s proxy (reward) is more problematic: "
    "reward conflates multiple behavioral components, and the direction of "
    "the proxy is ambiguous relative to the ELR construct. Future runs with "
    "corrected metric persistence would resolve these ambiguities."
)

heading("6.5 Series Trajectory", level=2)
body(
    "Protocol 6 adds one new positive finding to the series: emergent constraint "
    "fields create mechanistically distinct temporal coupling. It also extends two "
    "established patterns: behavioral homogenization is the dominant attractor "
    "regardless of constraint origin, and constraint overhead reduces reward "
    "relative to unconstrained baselines regardless of constraint form. The "
    "cumulative message of P2\u2013P6 is that alignment-relevant behavioral "
    "properties do not emerge from constraint architecture alone in this "
    "three-agent configuration, whether imposed or emergent. Whether larger "
    "populations, more heterogeneous architectures, or different reward "
    "structures would break the homogenization attractor remains an open question "
    "the current series cannot answer."
)

# ── 7. CONCLUSION ─────────────────────────────────────────────────────────────
heading("7. Conclusion")
body(
    "Protocol 6 tested whether changing the origin of a constraint \u2014 from "
    "externally imposed to emergent through agent co-constitution \u2014 changes "
    "alignment-relevant behavioral outcomes in a three-agent MARL system. The "
    "answer is mechanistically yes and behaviorally no. The emergent constraint "
    "field is causally active: it creates a robust, consistent temporal coupling "
    "between field structure and behavioral coordination (median r = \u22120.680 "
    "across 50 seeds, p < 0.001) that is entirely absent in the fixed-constraint "
    "and unconstrained conditions. This is a new finding in the series. But this "
    "mechanistic difference does not translate to better behavioral outcomes: the "
    "primary H2 comparison (emergent vs. fixed constraint) does not reach "
    "significance (p = 0.069), the H3 reversal contradicts the theoretical "
    "framing of local field perception as a source of heterogeneity, and behavioral "
    "homogenization remains the dominant attractor across all constrained "
    "conditions. The series pattern continues: imposing, deepening, or "
    "internalizing constraints does not produce sustained alignment-relevant "
    "behavior in this architecture. Passive emergence, like passive imposition, "
    "is insufficient as a governance strategy."
)

# ── INTEGRITY CHAIN ───────────────────────────────────────────────────────────
heading("Integrity Chain")

doc.add_paragraph()
tbl5 = doc.add_table(rows=1, cols=3)
tbl5.style = 'Table Grid'
for cell, text in zip(tbl5.rows[0].cells, ["Event", "Commit / Tag", "Date"]):
    cell.text = ''
    r = cell.paragraphs[0].add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)

integrity_rows = [
    ("Project initialized",                  "9765ac0 (tag: p6-init)",                 "2026-03-27"),
    ("P5 engine ported",                      "8d62d0a",                                "2026-03-27"),
    ("Deviation 1 \u2014 signal weight scale","adca6fc",                                "2026-03-28"),
    ("Pilot config + run script",             "326d9f4",                                "2026-03-28"),
    ("Pilot complete (60 runs)",              "8bf2f6f",                                "2026-03-28"),
    ("Preregistration SHA-256 locked",        "1ea0b51 (tag: p6-prereg-lock)",          "2026-03-28"),
    ("Zenodo DOI locked",                     "ff3fdb6",                                "2026-03-28"),
    ("AI committee predictions committed",    "7625fe1",                                "2026-03-28"),
    ("Four-condition confirmatory harness",   "da5cc47 (tag: p6-confirmatory-pre-run)", "2026-03-29"),
    ("Confirmatory runs complete (200 seeds)","4a12315 (tag: p6-confirmatory-complete)","2026-03-30"),
    ("Statistical analysis added",            "829be57",                                "2026-03-30"),
    ("Confirmatory findings report",          "dff15e6",                                "2026-03-30"),
    ("AI committee scoring report",           "eb343b5",                                "2026-03-30"),
    ("SSS/ELR metric persistence fix",        "b1c1bb1",                                "2026-03-30"),
]
for rd in integrity_rows:
    add_table_row(tbl5, rd)

doc.add_paragraph()

body(
    "Preregistration PDF: docs/Protocol6_Preregistration_v3.pdf\n"
    "SHA-256: c286a89037966e56630f5d3ce4cdb4a1621ce7ce9a0f7e7d477a06b8495dffc3\n"
    "Preregistration DOI: 10.5281/zenodo.19297509\n\n"
    "Logged deviations:\n"
    "  Deviation 1 (adca6fc): Signal weights scaled 0.1\u00d7 pre-pilot "
    "(DECLARE 0.3\u21920.03, QUERY 0.1\u21920.01, RESPOND 0.2\u21920.02).\n"
    "  Deviation 2 (post-pilot): Field formation criterion revised from "
    "field_std > 0.05 sustained >50 epochs to field_mean > 0.01 (activity-based). "
    "Research question revised from spatial differentiation to temporal coupling."
)

# ── REFERENCES ────────────────────────────────────────────────────────────────
heading("References")

refs = [
    "Axelrod, R. (1986). An evolutionary approach to norms. "
    "American Political Science Review, 80(4), 1095\u20131111.",

    "Goldbeter, A. (1996). Biochemical Oscillations and Cellular Rhythms. "
    "Cambridge University Press.",

    "Kessin, R. H. (2001). Dictyostelium: Evolution, Cell Biology, and "
    "the Development of Multicellularity. Cambridge University Press.",

    "Ostrom, E. (1990). Governing the Commons: The Evolution of Institutions "
    "for Collective Action. Cambridge University Press.",

    "Tisler, B. (2026a). Ethics as Structural Necessity in Multi-Agent "
    "Reinforcement Learning Systems: Protocol 2 — Fixed Ethical Tax and "
    "the Query-Flooding Attractor [Preregistration]. Zenodo. "
    "https://doi.org/10.5281/zenodo.18929040",

    "Tisler, B. (2026b). Protocol 3 Preregistration: Constraint-Ethics-Necessity "
    "— Enforcement Opacity and the Limits of Regulatory Constraint Design "
    "[Preregistration]. Zenodo. https://doi.org/10.5281/zenodo.19096602",

    "Tisler, B. (2026c). Protocol 4 Preregistration: Ethics as Emergent "
    "Constraint Response \u2014 From Mimesis to Phase Transition in Multi-Agent "
    "Systems [Preregistration]. Zenodo. https://doi.org/10.5281/zenodo.19005417",

    "Tisler, B. (2026d). Protocol 5 Preregistration: Ethics as Emergent "
    "Constraint Response \u2014 Temporal Integration Span and Prosocial Constraint "
    "Architecture as Necessary Conditions for Ethical Convergence "
    "[Preregistration]. Zenodo. https://doi.org/10.5281/zenodo.19038790",

    "Tisler, B. (2026e). Protocol 6 Preregistration: Emergent Constraint "
    "Landscapes as a Structural Alternative to Imposed Regulatory Constraint "
    "in Multi-Agent Systems [Preregistration]. Zenodo. "
    "https://doi.org/10.5281/zenodo.19297509",
]

for ref in refs:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent   = Inches(0.4)
    p.paragraph_format.first_line_indent = Inches(-0.4)
    p.paragraph_format.space_after   = Pt(4)
    r = p.add_run(ref)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

# ── SAVE ──────────────────────────────────────────────────────────────────────
out_path = r"D:/Claude_Code/project_12/Protocol6_Results_Paper.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
